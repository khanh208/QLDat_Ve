from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BAO_CAO_DIR = ROOT / "Bao_Cao"


def get_main_report_path() -> Path:
    exact = BAO_CAO_DIR / "báo cáo.docx"
    if exact.exists():
        return exact
    return next(
        p for p in BAO_CAO_DIR.iterdir()
        if p.suffix.lower() == ".docx" and p.name == "báo cáo.docx"
    )


def find_paragraph_index(doc: Document, text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise ValueError(f"Không tìm thấy đoạn: {text}")


def find_first_matching_paragraph_index(doc: Document, candidates: list[str]) -> int:
    for candidate in candidates:
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == candidate:
                return idx
    raise ValueError(f"Không tìm thấy đoạn nào trong danh sách: {candidates}")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_before(anchor, text: str, style: str = "Normal"):
    paragraph = anchor.insert_paragraph_before(text)
    paragraph.style = style
    return paragraph


def main() -> None:
    report_path = get_main_report_path()
    backup_path = report_path.with_name(report_path.stem + "_backup_before_chapter5.docx")
    if not backup_path.exists():
        shutil.copy2(report_path, backup_path)

    doc = Document(str(report_path))

    chapter5_idx = find_first_matching_paragraph_index(doc, [
        "CHƯƠNG 5: TEST - KIỂM THỬ",
        "PHẦN 5: THỰC THI KIỂM THỬ (TEST EXECUTION)",
    ])
    chapter6_idx = find_paragraph_index(doc, "CHƯƠNG 6: TRIỂN KHAI CÀI ĐẶT")

    # Xóa phần giữa Chương 5 và Chương 6 để chèn lại sạch sẽ khi chạy lại script.
    while chapter6_idx - chapter5_idx > 1:
        remove_paragraph(doc.paragraphs[chapter5_idx + 1])
        chapter6_idx -= 1

    doc.paragraphs[chapter5_idx].text = "PHẦN 5: THỰC THI KIỂM THỬ (TEST EXECUTION)"
    doc.paragraphs[chapter5_idx].style = "Heading 1"

    chapter6_anchor = doc.paragraphs[chapter6_idx]

    content = [
        ("5.1. Kiểm thử chức năng (Functional Testing)", "Heading 2"),
        (
            "Trong phạm vi báo cáo này, kiểm thử chức năng được thực hiện tập trung trên module BookingService của hệ thống QLDatVe. Bộ FunctionalTestingSuite gọi lại các nhóm test BookingValidationTest, BookingSafetyTest, BookingLifecycleTest và BookingValidationGapsTest để kiểm tra đồng thời các luồng nghiệp vụ chính, các ca dữ liệu không hợp lệ, các trường hợp race condition và vòng đời booking.",
            "Normal",
        ),
        (
            "Kết quả thực thi hiện tại cho thấy có 30 test case chức năng đã được chạy, trong đó 22 test đạt và 8 test không đạt. Tám test fail này phản ánh đúng các lỗi đang tồn tại trong implementation hiện tại, gồm 3 lỗi liên quan đến Safety và race condition khi nhiều yêu cầu cùng giữ một ghế, cùng với 5 lỗi validation như paymentMethod null, paymentMethod không hợp lệ, danh sách ghế chứa null hoặc chuỗi rỗng vẫn chưa bị chặn đúng theo mong đợi nghiệp vụ.",
            "Normal",
        ),
        (
            "Minh chứng cho mục này được lưu trong log 5.1_Functional_Testing.log và bộ 30 test case gốc tại TestCase.xlsx.",
            "Normal",
        ),
        ("5.2. Kiểm thử tích hợp (Integration Testing)", "Heading 2"),
        (
            "Kiểm thử tích hợp được bổ sung bằng BookingApiIntegrationTest với WebMvcTest, MockMvc và @MockBean. Ở bản rút gọn cho báo cáo, phần này giữ hai test case cơ bản: tạo booking CASH thành công và trả lỗi khi service phát hiện seat conflict.",
            "Normal",
        ),
        (
            "Kết quả thực thi cho nhóm này là 2 trên 2 test pass. Nhóm test tích hợp giúp chứng minh controller đang trả mã HTTP và payload đúng với kịch bản nghiệp vụ đã được mock.",
            "Normal",
        ),
        (
            "Các test case tích hợp cơ bản đã được bổ sung vào workbook TestCase.xlsx tại TC-31 và TC-32.",
            "Normal",
        ),
        ("5.3. Kiểm thử hệ thống (System Testing)", "Heading 2"),
        (
            "Do phạm vi bài làm tập trung vào backend, kiểm thử hệ thống được mô phỏng ở mức MockMvc flow. Bộ BookingSystemFlowMockMvcTest kiểm tra hai hành trình chính: luồng CASH từ tra cứu chuyến đến tạo booking và xem My Bookings, cùng với luồng MOMO từ tạo booking PENDING đến checkout, confirm callback và kiểm tra lịch sử booking.",
            "Normal",
        ),
        (
            "Kết quả thực thi cho nhóm này là 2 trên 2 test pass. Dù chưa phải system testing end-to-end với môi trường thật, phần kiểm thử này vẫn có giá trị trong việc kiểm tra chuỗi hành động gần với hành vi vận hành thực tế của hệ thống hơn so với việc chỉ kiểm từng endpoint đơn lẻ.",
            "Normal",
        ),
        (
            "Các test case hệ thống đã được bổ sung vào TestCase.xlsx tại TC-33 và TC-34.",
            "Normal",
        ),
        ("5.4. Kiểm thử hiệu năng (Performance Testing)", "Heading 2"),
        (
            "Kiểm thử hiệu năng được thực hiện ở mức smoke test bằng BookingPerformanceSmokeTest trong môi trường JUnit + Mockito. Test thứ nhất đo độ trễ của createBooking trong 150 lượt gọi liên tiếp. Test thứ hai đo thời gian hoàn tất của 12 vòng đặt vé đồng thời với 4 luồng mỗi vòng.",
            "Normal",
        ),
        (
            "Kết quả đo hiện tại cho thấy createBooking có Average = 2.621 ms, P95 = 4.290 ms, Max = 5.081 ms. Với kịch bản đồng thời, Average = 32.751 ms, P95 = 129.778 ms và Max = 129.778 ms. Cả hai test đều đạt ngưỡng kiểm thử đã đặt ra trong mã nguồn, vì vậy nhóm performance smoke test được ghi nhận là pass.",
            "Normal",
        ),
        (
            "Các test case hiệu năng đã được bổ sung vào TestCase.xlsx tại TC-35 và TC-36; số liệu chi tiết được lưu ở target/performance-reports.",
            "Normal",
        ),
        ("5.5. Kiểm thử bảo mật cơ bản (Security Testing)", "Heading 2"),
        (
            "Kiểm thử bảo mật cơ bản được bổ sung bằng SecurityAccessIntegrationTest, tập trung vào xác thực và phân quyền trên các endpoint quan trọng. Ở bản rút gọn cho báo cáo, phần này giữ ba test case cơ bản: cho phép anonymous truy cập /api/trips, chặn anonymous ở /api/bookings/my-bookings, và cho phép admin truy cập /api/admin/bookings.",
            "Normal",
        ),
        (
            "Kết quả thực thi nhóm này là 3 trên 3 test pass. Các ca kiểm thử bảo mật cơ bản đã được bổ sung vào TestCase.xlsx từ TC-37 đến TC-39, đồng thời được tổng hợp trong sheet Chuong5_JUnitMock của TestReport.xlsx.",
            "Normal",
        ),
        ("5.6. Kiểm thử hồi quy (Regression Testing)", "Heading 2"),
        (
            "Kiểm thử hồi quy được triển khai bằng RegressionStableSuite. Suite này sử dụng SuiteAssertions để chạy lại các nhóm test ổn định gồm BookingValidationTest, BookingLifecycleTest và SecurityAccessIntegrationTest sau khi bổ sung các test mới cho Chương 5. Mục tiêu là bảo đảm các nhánh nghiệp vụ ổn định trước đó không bị ảnh hưởng bởi phần mở rộng kiểm thử mới.",
            "Normal",
        ),
        (
            "Kết quả thực thi là 1 trên 1 regression suite pass. Test case hồi quy tương ứng đã được bổ sung vào TestCase.xlsx tại TC-40.",
            "Normal",
        ),
        ("5.7. Thống kê kết quả", "Heading 2"),
        (
            "Tổng hợp Chương 5 hiện có: Functional Testing 30 tested, 22 passed, 8 failed; Integration Testing 2 tested, 2 passed; System Testing 2 tested, 2 passed; Performance Testing 2 tested, 2 passed; Security Testing 3 tested, 3 passed; Regression Testing 1 suite, 1 passed.",
            "Normal",
        ),
        (
            "Nếu tính toàn bộ các mục đã thực thi trong Chương 5, có 40 lượt thực thi được ghi nhận, trong đó 32 pass và 8 fail, tương đương tỷ lệ pass khoảng 80.00%. Lưu ý rằng 8 fail hiện tại chỉ nằm ở nhóm functional và phản ánh các bug thật của implementation BookingService, không phải lỗi của môi trường test.",
            "Normal",
        ),
        (
            "Các số liệu tổng hợp cho Chương 5 đã được thêm vào TestReport.xlsx ở các sheet Chuong5_JUnitMock, Chuong5_Mapping và Chuong5_Artifacts để phục vụ chèn bảng minh chứng vào Word.",
            "Normal",
        ),
    ]

    for text, style in content:
        insert_paragraph_before(chapter6_anchor, text, style)

    saved_to = report_path
    try:
        doc.save(str(report_path))
    except PermissionError:
        saved_to = report_path.with_name("báo cáo_rut_gon_chapter5.docx")
        doc.save(str(saved_to))
    print(saved_to)
    print(backup_path)


if __name__ == "__main__":
    main()
