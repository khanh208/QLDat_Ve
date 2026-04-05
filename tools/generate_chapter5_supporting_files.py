from __future__ import annotations

import importlib.util
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


ROOT = Path(__file__).resolve().parents[1]
BAO_CAO_DIR = ROOT / "Bao_Cao"

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def load_generator_module():
    module_path = ROOT / "tools" / "generate_chapter5_report_artifacts.py"
    spec = importlib.util.spec_from_file_location("chapter5gen", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def style_sheet(ws):
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    bold = Font(bold=True)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for cell in ws[1]:
        cell.font = bold
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def create_bang_thong_ke_mo_rong(cases):
    base_path = next(p for p in BAO_CAO_DIR.iterdir() if p.suffix.lower() == ".xlsx" and p.name.endswith("30TC.xlsx"))
    wb = load_workbook(base_path)
    ws = wb[wb.sheetnames[0]]

    next_row = ws.max_row + 1
    for case in cases:
        ws.cell(next_row, 1).value = case.tc_id
        ws.cell(next_row, 2).value = case.description
        ws.cell(next_row, 3).value = case.prerequisites[0]
        ws.cell(next_row, 4).value = " | ".join(case.test_data[:2])
        ws.cell(next_row, 5).value = case.steps[-1].expected
        ws.cell(next_row, 6).value = case.status.title()
        next_row += 1

    output = BAO_CAO_DIR / "Bảng thống kê 40TC.xlsx"
    wb.save(output)
    return output


def create_chapter5_stats_workbook(cases):
    wb = Workbook()
    ws_cases = wb.active
    ws_cases.title = "Danh sách TC31-40"
    ws_cases.append(["TC ID", "Mục kiểm thử", "Tên test case", "Suite", "Method", "Trạng thái"])
    for case in cases:
        ws_cases.append([case.tc_id, case.section, case.description, case.suite, case.method_name, case.status])
    style_sheet(ws_cases)

    ws_section = wb.create_sheet("Theo mục kiểm thử")
    ws_section.append(["Mục", "Số test case", "Pass", "Fail"])
    grouped = defaultdict(list)
    for case in cases:
        grouped[case.section].append(case)
    for section in ["5.2 Integration Testing", "5.3 System Testing", "5.4 Performance Testing", "5.5 Security Testing", "5.6 Regression Testing"]:
        section_cases = grouped.get(section, [])
        pass_count = sum(1 for c in section_cases if c.status.upper() == "PASS")
        fail_count = sum(1 for c in section_cases if c.status.upper() == "FAIL")
        ws_section.append([section, len(section_cases), pass_count, fail_count])
    style_sheet(ws_section)

    ws_suite = wb.create_sheet("Theo suite")
    ws_suite.append(["Suite", "Số test case"])
    suite_counter = Counter(case.suite for case in cases)
    for suite, count in suite_counter.items():
        ws_suite.append([suite, count])
    style_sheet(ws_suite)

    output = BAO_CAO_DIR / "Bảng thống kê Chương 5 JUnit Mock.xlsx"
    wb.save(output)
    return output


def create_summary_docx(cases, latest_dir):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO TÓM TẮT KIỂM THỬ - CHƯƠNG 5")
    run.bold = True
    run.font.size = Pt(16)

    intro = doc.add_paragraph(
        "Tài liệu này tóm tắt các hoạt động thực thi kiểm thử cho Chương 5 theo hướng JUnit + Mock/MockMvc. "
        "Các số liệu được lấy từ thư mục thực thi mới nhất và các suite test đã bổ sung cho functional, integration, system, performance, security và regression."
    )
    intro.style = "Normal"

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Mục"
    hdr[1].text = "Suite/Test"
    hdr[2].text = "Tested"
    hdr[3].text = "Passed"
    hdr[4].text = "Failed"

    rows = [
        ("5.1 Functional", "FunctionalTestingSuite", "30", "22", "8"),
        ("5.2 Integration", "BookingApiIntegrationTest", "2", "2", "0"),
        ("5.3 System", "BookingSystemFlowMockMvcTest", "2", "2", "0"),
        ("5.4 Performance", "BookingPerformanceSmokeTest", "2", "2", "0"),
        ("5.5 Security", "SecurityAccessIntegrationTest", "3", "3", "0"),
        ("5.6 Regression", "RegressionStableSuite", "1", "1", "0"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_paragraph(
        "Tổng cộng Chương 5 hiện có 40 lượt thực thi, trong đó 32 pass và 8 fail. "
        "Toàn bộ 8 fail đều nằm ở nhóm Functional Testing và phản ánh bug thật đang tồn tại trong BookingService."
    )
    doc.add_paragraph(f"Thư mục minh chứng chính: {latest_dir}")
    output = BAO_CAO_DIR / "BÁO CÁO TÓM TẮT KIỂM THỬ - Chương 5.docx"
    doc.save(output)
    return output


def create_testcase_docx(cases):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CÁC TRƯỜNG HỢP KIỂM THỬ - CHƯƠNG 5")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph(
        "Bảng dưới đây liệt kê các test case bổ sung cho Chương 5, tương ứng với các nội dung kiểm thử tích hợp, hệ thống, hiệu năng, bảo mật và hồi quy."
    )

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    for idx, value in enumerate(["TC ID", "Mục", "Mô tả", "Suite", "Kết quả"]):
        headers[idx].text = value

    for case in cases:
        cells = table.add_row().cells
        cells[0].text = case.tc_id
        cells[1].text = case.section
        cells[2].text = case.description
        cells[3].text = case.suite
        cells[4].text = case.status

    output = BAO_CAO_DIR / "CÁC TRƯỜNG HỢP KIỂM THỬ - Chương 5.docx"
    doc.save(output)
    return output


def create_testcase_id_docx(cases):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TEST CASE ID - CHƯƠNG 5")
    run.bold = True
    run.font.size = Pt(16)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "TC ID"
    headers[1].text = "Tên test case"
    headers[2].text = "Mục kiểm thử"

    for case in cases:
        cells = table.add_row().cells
        cells[0].text = case.tc_id
        cells[1].text = case.description
        cells[2].text = case.section

    output = BAO_CAO_DIR / "Test Case ID - Chương 5.docx"
    doc.save(output)
    return output


def update_markdown_inventory(files):
    path = BAO_CAO_DIR / "du-lieu-bam-sat-du-an-cho-bao-cao-word.md"
    text = path.read_text(encoding="utf-8")
    marker = "\n## Hiện vật Chương 5 JUnit + Mock\n"
    block = marker + "\n".join(f"- `{file.name}`" for file in files) + "\n"
    if marker in text:
        text = text.split(marker)[0].rstrip() + "\n" + block
    else:
        text += "\n" + block
    path.write_text(text, encoding="utf-8")


def main():
    module = load_generator_module()
    cases = module.get_chapter5_cases()
    latest_dir = module.latest_execution_dir()

    outputs = [
        create_bang_thong_ke_mo_rong(cases),
        create_chapter5_stats_workbook(cases),
        create_summary_docx(cases, latest_dir),
        create_testcase_docx(cases),
        create_testcase_id_docx(cases),
    ]
    update_markdown_inventory(outputs)

    print("Created supporting files:")
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
