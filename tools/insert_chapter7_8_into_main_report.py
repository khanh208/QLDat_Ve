from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

from generate_chapter7_8_supporting_files import CHAPTER7_TEXT, CHAPTER8_TEXT


ROOT = Path(__file__).resolve().parents[1]
BAO_CAO_DIR = ROOT / "Bao_Cao"


def get_main_report_path() -> Path:
    for path in BAO_CAO_DIR.iterdir():
        if path.suffix.lower() == ".docx" and path.name == "báo cáo.docx":
            return path
    raise FileNotFoundError("Không tìm thấy file báo cáo chính.")


def find_index_startswith(doc: Document, prefix: str) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return idx
    return None


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main() -> None:
    report_path = get_main_report_path()
    backup_path = report_path.with_name(report_path.stem + "_backup_before_chapter7_8.docx")
    if not backup_path.exists():
        shutil.copy2(report_path, backup_path)

    document = Document(str(report_path))

    chapter7_idx = find_index_startswith(document, "CHƯƠNG 7:")
    references_idx = find_index_startswith(document, "TÀI LIỆU THAM KHẢO")

    if chapter7_idx is None:
        raise ValueError("Không tìm thấy vị trí bắt đầu Chương 7.")
    if references_idx is None:
        raise ValueError("Không tìm thấy mục TÀI LIỆU THAM KHẢO.")

    while references_idx - chapter7_idx > 1:
        remove_paragraph(document.paragraphs[chapter7_idx + 1])
        references_idx -= 1

    document.paragraphs[chapter7_idx].text = "CHƯƠNG 7: BÁO CÁO LỖI (BUG REPORT)"
    document.paragraphs[chapter7_idx].style = "Heading 1"

    anchor = document.paragraphs[references_idx]
    for text, style in CHAPTER7_TEXT[1:] + CHAPTER8_TEXT:
        paragraph = anchor.insert_paragraph_before(text)
        paragraph.style = style

    try:
        document.save(str(report_path))
        print(report_path)
    except PermissionError:
        fallback = report_path.with_name(report_path.stem + "_chapter7_8_updated.docx")
        document.save(str(fallback))
        print(fallback)
    print(backup_path)


if __name__ == "__main__":
    main()
