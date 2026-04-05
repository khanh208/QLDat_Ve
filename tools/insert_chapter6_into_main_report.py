from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

from generate_chapter6_supporting_files import CHAPTER6_TEXT


ROOT = Path(__file__).resolve().parents[1]
BAO_CAO_DIR = ROOT / "Bao_Cao"


def get_main_report_path() -> Path:
    preferred = [p for p in BAO_CAO_DIR.iterdir() if p.suffix.lower() == ".docx" and p.name == "báo cáo.docx"]
    if preferred:
        return preferred[0]
    for path in BAO_CAO_DIR.iterdir():
        if path.suffix.lower() == ".docx" and "báo cáo" in path.name.lower() and "backup" not in path.name.lower() and "rút_gọn" not in path.name.lower():
            return path
    raise FileNotFoundError("Không tìm thấy file báo cáo chính.")


def find_index_exact(doc: Document, text: str) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    return None


def find_index_startswith(doc: Document, prefix: str) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return idx
    return None


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main() -> None:
    report_path = get_main_report_path()
    backup_path = report_path.with_name(report_path.stem + "_backup_before_chapter6.docx")
    if not backup_path.exists():
        shutil.copy2(report_path, backup_path)

    document = Document(str(report_path))

    chapter6_idx = find_index_exact(document, "CHƯƠNG 6: TRIỂN KHAI CÀI ĐẶT")
    if chapter6_idx is None:
        chapter6_idx = find_index_exact(document, "CHƯƠNG 6: KIỂM THỬ TỰ ĐỘNG (AUTOMATION TESTING)")
    if chapter6_idx is None:
        raise ValueError("Không tìm thấy vị trí bắt đầu Chương 6 trong báo cáo.")

    chapter7_idx = find_index_startswith(document, "CHƯƠNG 7:")
    if chapter7_idx is None:
        raise ValueError("Không tìm thấy vị trí bắt đầu Chương 7 trong báo cáo.")

    while chapter7_idx - chapter6_idx > 1:
        remove_paragraph(document.paragraphs[chapter6_idx + 1])
        chapter7_idx -= 1

    document.paragraphs[chapter6_idx].text = "CHƯƠNG 6: KIỂM THỬ TỰ ĐỘNG (AUTOMATION TESTING)"
    document.paragraphs[chapter6_idx].style = "Heading 1"

    anchor = document.paragraphs[chapter7_idx]
    for text, style in CHAPTER6_TEXT[1:]:
        paragraph = anchor.insert_paragraph_before(text)
        paragraph.style = style

    try:
        document.save(str(report_path))
        print(report_path)
    except PermissionError:
        fallback = report_path.with_name(report_path.stem + "_chapter6_updated.docx")
        document.save(str(fallback))
        print(fallback)
    print(backup_path)


if __name__ == "__main__":
    main()
